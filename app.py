import streamlit as st
import json
import io
import re
import zipfile
import base64
import urllib.parse
import unicodedata
import docx
from docx import Document
from docx.shared import Inches

# --- TERMINAL-LOGGNING ---
def log(msg):
    print(f"[LOGG] {msg}")

# --- HJÄLPFUNKTIONER ---

def add_hyperlink(paragraph, text, url):
    """Skapar en klickbar blå länk i Word-dokumentet."""
    try:
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), "0000FF")
        rPr.append(c)
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)
        new_run.append(rPr)
        t = docx.oxml.shared.OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception:
        paragraph.add_run(f"{text} ({url})")

def normalize(name):
    if not name: return ""
    n = unicodedata.normalize('NFD', name)
    n = "".join([c for c in n if unicodedata.category(c) != 'Mn']).lower()
    n = n.replace('\xa0', ' ').replace('□', ' ')
    return re.sub(r'[^a-z0-9]', '', n)

def find_image_in_uploads(target_name, target_id, uploaded_images):
    t_name_norm = normalize(target_name)
    t_id_str = str(target_id) if target_id else None
    for path, file_obj in uploaded_images.items():
        if t_id_str and t_id_str in path: return file_obj
        if t_name_norm and t_name_norm in normalize(path): return file_obj
    return None

def extract_unique_legacy_images(pool):
    unique_images = {}
    if not pool or 'numToAttrib' not in pool: return []
    for attr in pool['numToAttrib'].values():
        val = attr[0] if isinstance(attr, list) else str(attr)
        if val.startswith('image-'):
            try:
                parts = val.split('-')
                if len(parts) >= 3:
                    raw = base64.b64decode(urllib.parse.unquote(parts[-1]))
                    info = json.loads(urllib.parse.unquote(raw.decode('utf-8')))
                    fname = info.get('fileName', 'image.png')
                    fid = info.get('boxFileId') or info.get('fileId')
                    name_key = normalize(fname)
                    if name_key not in unique_images:
                        unique_images[name_key] = {'name': fname, 'id': fid, 'link': info.get('boxSharedLink')}
            except Exception: pass
    return list(unique_images.values())

def process_node_list(nodes, parent_obj, current_p, uploaded_images, warnings, is_list=False, list_type='bullet'):
    """Hanterar moderna noder inkl. tabeller, listor och bilder."""
    missing_any = False
    for node in nodes:
        n_type = node.get('type')
        
        if n_type == 'text':
            text = node.get('text', '')
            marks = node.get('marks', [])
            link = next((m['attrs']['href'] for m in marks if m['type'] == 'link'), None)
            if link:
                add_hyperlink(current_p, text, link)
            else:
                run = current_p.add_run(text)
                for mark in marks:
                    m_type = mark.get('type')
                    if m_type == 'strong': run.bold = True
                    if m_type == 'italic': run.italic = True
                    if m_type == 'underline': run.underline = True

        elif n_type == 'hard_break':
            if current_p: current_p.add_run('\n')

        elif n_type == 'paragraph':
            style = 'List Bullet' if is_list and list_type == 'bullet' else None
            if is_list and list_type == 'number': style = 'List Number'
            # Skapa nytt stycke i föräldern (Document eller TableCell)
            p = parent_obj.add_paragraph(style=style)
            if 'content' in node:
                if process_node_list(node['content'], parent_obj, p, uploaded_images, warnings):
                    missing_any = True

        elif n_type == 'heading':
            level = node.get('attrs', {}).get('level', 1)
            # TableCells stöder inte add_heading, använd fetstil istället
            if hasattr(parent_obj, 'add_heading'):
                p = parent_obj.add_heading('', level=level)
            else:
                p = parent_obj.add_paragraph()
                p.bold = True
            if 'content' in node:
                if process_node_list(node['content'], parent_obj, p, uploaded_images, warnings):
                    missing_any = True

        elif n_type in ['bullet_list', 'ordered_list']:
            l_type = 'bullet' if n_type == 'bullet_list' else 'number'
            for item in node.get('content', []):
                if 'content' in item:
                    if process_node_list(item['content'], parent_obj, None, uploaded_images, warnings, is_list=True, list_type=l_type):
                        missing_any = True

        elif n_type == 'table':
            rows_data = node.get('content', [])
            if not rows_data: continue
            num_cols = len(rows_data[0].get('content', []))
            table = parent_obj.add_table(rows=len(rows_data), cols=num_cols)
            table.style = 'Table Grid'
            for r_idx, row_node in enumerate(rows_data):
                cells_data = row_node.get('content', [])
                for c_idx, cell_node in enumerate(cells_data):
                    if c_idx < num_cols:
                        cell = table.rows[r_idx].cells[c_idx]
                        # Töm cellens första standardparagraf om vi lägger till egna
                        if 'content' in cell_node:
                            # För att undvika en tom rad i början av cellen, 
                            # kan vi passera cellen som parent för dess innehåll
                            process_node_list(cell_node['content'], cell, None, uploaded_images, warnings)

        elif n_type == 'image':
            attrs = node.get('attrs', {})
            img = find_image_in_uploads(attrs.get('fileName'), attrs.get('boxFileId'), uploaded_images)
            if img:
                # Bilden läggs i ett nytt stycke om inget current_p finns
                target = current_p if current_p else parent_obj.add_paragraph()
                target.add_run().add_picture(io.BytesIO(img.getvalue()), width=Inches(5.0))
            else:
                missing_any = True
                target = current_p if current_p else parent_obj.add_paragraph()
                target.add_run(f"\n⚠️ [BILD SAKNAS: {attrs.get('fileName')}]\n").italic = True
                
    return missing_any

# --- APP ---
st.set_page_config(page_title="BoxNote Pro", layout="centered")
st.title("📝 BoxNote Pro-konverterare")

uploaded_files = st.file_uploader("Dra in .boxnote och bilder här", accept_multiple_files=True)

if uploaded_files:
    notes = [f for f in uploaded_files if f.name.endswith('.boxnote')]
    imgs = {f.name: f for f in uploaded_files if f.name.lower().endswith(('.png', '.jpg', '.jpeg'))}
    
    if notes and st.button(f"Konvertera {len(notes)} noter"):
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "a", zipfile.ZIP_DEFLATED) as zip_file:
            for n_file in notes:
                try:
                    data = json.loads(n_file.getvalue().decode("utf-8"))
                    doc = Document()
                    missing_images = False
                    if 'doc' in data:
                        missing_images = process_node_list(data['doc'].get('content', []), doc, None, imgs, [])
                    elif 'atext' in data:
                        doc.add_heading(n_file.name.replace(".boxnote", ""), 0)
                        for line in data['atext'].get('text', '').split('\n'):
                            doc.add_paragraph(line)
                        legacy_imgs = extract_unique_legacy_images(data.get('pool', {}))
                        if legacy_imgs:
                            doc.add_heading("Bifogade bilder", level=2)
                            for info in legacy_imgs:
                                img_data = find_image_in_uploads(info['name'], info['id'], imgs)
                                if img_data:
                                    doc.add_picture(io.BytesIO(img_data.getvalue()), width=Inches(5.5))
                                else:
                                    missing_images = True
                                    p = doc.add_paragraph(f"⚠️ Bild saknas: {info['name']} - ")
                                    if info.get('link'): add_hyperlink(p, "SE BILD PÅ BOX", info['link'])
                    base_name = n_file.name.replace(".boxnote", ".docx")
                    final_name = f"[FIXA] {base_name}" if missing_images else base_name
                    d_io = io.BytesIO()
                    doc.save(d_io)
                    zip_file.writestr(final_name, d_io.getvalue())
                except Exception as e:
                    st.error(f"Fel vid {n_file.name}: {e}")
        st.success("Klar!")
        st.download_button("📥 Ladda ner ZIP", zip_buf.getvalue(), "box_export.zip")