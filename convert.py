import json
import os
import re
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def add_hyperlink(paragraph, text, url):
    """Hjälpfunktion för att lägga till en klickbar länk i ett stycke."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Ge länken den klassiska blå färgen och understrykning
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), "0000FF")
    rPr.append(c)
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    t = docx.oxml.shared.OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def find_image_fuzzy(specific_images_folder, target_name):
    if not os.path.isdir(specific_images_folder):
        return None
    def normalize(name):
        return re.sub(r'[^a-zA-Z0-9]', '', name.replace('\xa0', ' ')).lower()
    target_norm = normalize(target_name)
    try:
        all_files = os.listdir(specific_images_folder)
        for filename in all_files:
            if normalize(filename) == target_norm:
                return os.path.join(specific_images_folder, filename)
    except Exception:
        pass
    return None

def process_content(content_list, doc_obj, paragraph_obj, specific_images_folder):
    for item in content_list:
        node_type = item.get('type')
        if node_type == 'text':
            text = item.get('text', '')
            marks = item.get('marks', [])
            
            # Kolla om texten är en länk
            link_url = next((m['attrs']['href'] for m in marks if m['type'] == 'link'), None)
            is_bold = any(m['type'] == 'strong' for m in marks)

            if link_url:
                add_hyperlink(paragraph_obj, text, link_url)
            else:
                run = paragraph_obj.add_run(text)
                if is_bold:
                    run.bold = True
        elif node_type == 'hard_break':
            paragraph_obj.add_run('\n')
        elif node_type == 'image':
            img_filename = item.get('attrs', {}).get('fileName')
            if img_filename:
                img_path = find_image_fuzzy(specific_images_folder, img_filename)
                if img_path:
                    doc_obj.add_picture(img_path, width=Inches(6))

def convert_boxnote_to_docx(source_folder, root_images_folder, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for filename in os.listdir(source_folder):
        if filename.endswith(".boxnote"):
            file_path = os.path.join(source_folder, filename)
            note_name_raw = filename.replace(".boxnote", "")
            specific_images_folder = os.path.join(root_images_folder, f"{note_name_raw} Images")

            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                nodes = data.get('doc', {}).get('content', [])
                doc = Document()
                
                for node in nodes:
                    node_type = node.get('type')
                    if node_type == 'heading':
                        p = doc.add_heading('', level=node.get('attrs', {}).get('level', 1))
                        process_content(node.get('content', []), doc, p, specific_images_folder)
                    elif node_type == 'paragraph':
                        p = doc.add_paragraph()
                        if 'content' in node:
                            process_content(node['content'], doc, p, specific_images_folder)
                    elif node_type == 'bullet_list':
                        for item in node.get('content', []):
                            for sub_node in item.get('content', []):
                                p = doc.add_paragraph(style='List Bullet')
                                process_content(sub_node.get('content', []), doc, p, specific_images_folder)
                    elif node_type == 'table':
                        rows = node.get('content', [])
                        if rows:
                            table = doc.add_table(rows=0, cols=len(rows[0].get('content', [])))
                            table.style = 'Table Grid'
                            for row_node in rows:
                                row_cells = table.add_row().cells
                                for idx, cell_node in enumerate(row_node.get('content', [])):
                                    for sub_p in cell_node.get('content', []):
                                        process_content(sub_p.get('content', []), doc, row_cells[idx].paragraphs[0], specific_images_folder)

                doc.save(os.path.join(output_folder, f"{note_name_raw}.docx"))
                print(f"✅ Konverterat med länkar och bilder: {filename}")
            except Exception as e:
                print(f"❌ Fel vid {filename}: {e}")

# Inställningar
base_path = '/mnt/c/Users/backa/Desktop/boxtest'
convert_boxnote_to_docx(base_path, os.path.join(base_path, 'Box Notes Images'), os.path.join(base_path, 'Konverterat'))